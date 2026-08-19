from __future__ import annotations

import base64
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.shared import Inches
from PIL import Image, ImageDraw


def make_text_image() -> bytes:
    image = Image.new("RGB", (640, 180), "white")
    drawing = ImageDraw.Draw(image)
    drawing.rectangle((12, 12, 627, 167), outline="black", width=3)
    drawing.text((48, 66), "DATABASE PORT: 3306", fill="black")
    buffer = BytesIO()
    image.save(buffer, format="PNG")
    return buffer.getvalue()


def image_data_url() -> str:
    encoded = base64.b64encode(make_text_image()).decode("ascii")
    return f"data:image/png;base64,{encoded}"


def create_docx_fixture(path: Path) -> None:
    document = Document()
    document.add_heading("数据库配置", level=1)
    document.add_paragraph("数据库默认端口为3306。")

    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "参数"
    table.rows[0].cells[1].text = "说明"
    row = table.add_row().cells
    row[0].text = "port"
    row[1].text = "数据库端口"

    document.add_paragraph("下面截图展示默认端口：")
    document.add_picture(BytesIO(make_text_image()), width=Inches(4.8))
    document.add_paragraph("截图后的说明文字。")
    document.save(str(path))
